import os
import sys
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
import random
import re

INPUT_FILE = "NGAN_HANG_CAU_HOI.docx"

def load_questions(filename):
    doc = Document(filename)
    paragraphs = [p for p in doc.paragraphs if p.text.strip()]
    questions = []
    i = 0

    while i < len(paragraphs):
        text = paragraphs[i].text.strip()
        if text.startswith("Câu "):
            # Sử dụng regex để loại bỏ phần "Câu X." hoặc "Câu X:" một cách triệt để
            question_text = re.sub(r"^Câu\s+\d+[\.\:]\s*", "", text, flags=re.IGNORECASE).strip()
            current = {"question": question_text, "answers": []}
            for j in range(1, 5):
                if i + j >= len(paragraphs): break
                para = paragraphs[i + j]
                answer_text = para.text.strip()

                # ✅ Đáp án đúng là in đậm hoặc có highlight
                is_correct = any(run.bold or (run.font.highlight_color is not None) for run in para.runs)

                current["answers"].append({
                    "text": answer_text,
                    "is_correct": is_correct
                })
            has_correct = any(ans["is_correct"] for ans in current["answers"])
            if len(current["answers"]) == 4:
                if has_correct:
                    questions.append(current)
                else:
                    print(f"⚠️  Cảnh báo: Câu hỏi '{question_text[:50]}...' không có đáp án đúng (in đậm). Đã bỏ qua.")
            i += 5
        else:
            i += 1

    return questions

def generate_exam(questions, num_questions, seed=None):
    if seed:
        random.seed(seed)
    selected = random.sample(questions, num_questions)
    exam, answer_key = [], []

    for idx, q in enumerate(selected, 1):
        shuffled = q["answers"].copy()
        random.shuffle(shuffled)

        try:
            correct_index = next(i for i, ans in enumerate(shuffled) if ans["is_correct"])
            exam.append({
                "index": idx,
                "question": q["question"],
                "answers": shuffled,
                "correct_letter": chr(65 + correct_index)
            })
            answer_key.append(f"Câu {idx}: {exam[-1]['correct_letter']}")
        except StopIteration:
            print(f"❌ Lỗi: Câu hỏi '{q['question'][:50]}...' bị mất đáp án đúng khi tạo đề.")
            continue
    
    return exam, answer_key

def export_exam_to_word(exam_data, answer_lines, filename, code):
    doc = Document()
    
    # Cài đặt font Times New Roman cho toàn bộ văn bản
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    # Cài đặt font cho các ký tự đặc biệt/tiếng Việt
    style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    
    title = doc.add_heading(f'ĐỀ THI - MÃ ĐỀ {code:03}', level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    for item in exam_data:
        p = doc.add_paragraph()
        run = p.add_run(f"Câu {item['index']}. {item['question']}")
        run.bold = True
        run.font.size = Pt(12)

        for i, ans in enumerate(item["answers"]):
            doc.add_paragraph(f"   {chr(65 + i)}. {ans['text']}")

    doc.add_page_break()
    doc.add_heading('ĐÁP ÁN', level=2)
    for ans in answer_lines:
        doc.add_paragraph(ans)

    doc.save(filename)

def main():
    print("📂 Kiểm tra file ngân hàng câu hỏi...")

    if not os.path.exists(INPUT_FILE):
        script_dir = os.path.dirname(os.path.abspath(__file__))
        print(f"❌ Không tìm thấy file '{INPUT_FILE}' trong thư mục hiện tại.")
        print(f"👉 Vui lòng thêm file Word có tên 'NGAN_HANG_CAU_HOI.docx' vào thư mục sau:\n📂 {script_dir}")
        sys.exit(1)


    questions = load_questions(INPUT_FILE)
    total_questions = len(questions)
    print(f"✅ Đã xác định số câu hỏi trong ngân hàng: {total_questions}")

    while True:
        try:
            num_exams = int(input("📄 Bạn muốn tạo bao nhiêu đề thi? "))
            if num_exams <= 0:
                raise ValueError
            break
        except ValueError:
            print("⚠️ Vui lòng nhập một số nguyên dương!")

    while True:
        try:
            num_per_exam = int(input("✍️ Mỗi đề thi có bao nhiêu câu hỏi? "))
            if num_per_exam <= 0 or num_per_exam > total_questions:
                raise ValueError
            break
        except ValueError:
            print(f"⚠️ Vui lòng nhập một số < {total_questions}!")

    for i in range(1, num_exams + 1):
        exam, answers = generate_exam(questions, num_per_exam, seed=i)
        filename = f"DE_THI_MA_DE_{i:03}.docx"
        export_exam_to_word(exam, answers, filename, i)
        print(f"✅ Đã tạo: {filename}")

if __name__ == "__main__":
    main()
