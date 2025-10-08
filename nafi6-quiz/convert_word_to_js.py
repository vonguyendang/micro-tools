import docx
import json
import re

def clean_text(text):
    """Loại bỏ các ký tự không cần thiết và khoảng trắng thừa."""
    return text.strip()

def find_correct_answer_by_bold(cell):
    """
    Hàm này được thiết kế lại hoàn toàn.
    Nó sẽ tìm chính xác 'run' văn bản nào được in đậm,
    và lấy ký tự đáp án từ chính 'run' đó.
    """
    for p in cell.paragraphs:
        for run in p.runs:
            # Nếu một run được in đậm và có chứa text
            if run.bold and run.text.strip():
                # Tìm ký tự A, B, C, D ở đầu của run in đậm đó
                match = re.match(r'^\s*([A-D])\.', run.text.strip())
                if match:
                    return match.group(1) # Trả về ký tự tìm được (A, B, C, hoặc D)
    return ""

def convert_docx_to_js(docx_path, js_path):
    """
    Đọc file DOCX có cấu trúc 3 cột và xuất ra file data.js.
    Đáp án đúng được xác định bằng chữ in đậm.
    """
    try:
        document = docx.Document(docx_path)
    except Exception as e:
        print(f"Lỗi: Không thể mở file '{docx_path}'. Hãy chắc chắn file tồn tại và không bị lỗi.")
        return

    if not document.tables:
        print(f"Lỗi: Không tìm thấy bảng nào trong file '{docx_path}'.")
        return
        
    table = document.tables[0]
    questions_data = []
    
    print("Bắt đầu xử lý file Word với logic chính xác...")
    
    for i, row in enumerate(table.rows):
        if len(row.cells) < 3:
            continue

        try:
            id_text = clean_text(row.cells[0].text)
            question_text = clean_text(row.cells[1].text)
            answers_cell = row.cells[2]
            
            if not id_text.isdigit() or not question_text:
                continue

            question_id = int(id_text)
            
            answers_dict = {}
            answer_lines = [line.strip() for line in answers_cell.text.split('\n') if line.strip()]
            
            for line in answer_lines:
                match = re.match(r'^([A-D])\.\s*(.*)', line, re.DOTALL)
                if match:
                    key = match.group(1)
                    value = match.group(2).strip()
                    answers_dict[key] = value

            correct_answer = find_correct_answer_by_bold(answers_cell)

            # Nếu không tìm thấy đáp án in đậm, thông báo để người dùng biết
            if not correct_answer:
                print(f"Cảnh báo: Không tìm thấy đáp án in đậm cho câu hỏi số {question_id}.")

            question_obj = {
                "id": question_id,
                "question": question_text,
                "answers": {
                    "A": answers_dict.get("A", ""),
                    "B": answers_dict.get("B", ""),
                    "C": answers_dict.get("C", ""),
                    "D": answers_dict.get("D", "")
                },
                "correctAnswer": correct_answer
            }
            questions_data.append(question_obj)

        except Exception as e:
            print(f"Lỗi khi xử lý hàng {i+1}: {e}")

    if not questions_data:
        print("Không có dữ liệu câu hỏi nào được trích xuất.")
        return

    with open(js_path, 'w', encoding='utf-8') as f:
        f.write("// data.js\n\n")
        f.write("const originalQuestions = ")
        json_string = json.dumps(questions_data, ensure_ascii=False, indent=4)
        f.write(json_string)
        f.write(";\n\n")
        f.write("const NUM_QUESTIONS = 70; // 70 câu hỏi ngẫu nhiên\n")
        f.write("const QUIZ_DURATION = 40 * 60; // 40 phút = 2400 giây\n")
    
    print(f"---")
    print(f"✅ Hoàn thành! Đã chuyển đổi {len(questions_data)} câu hỏi từ '{docx_path}' sang '{js_path}'.")

# --- THỰC THI SCRIPT ---
if __name__ == "__main__":
    input_file = "NganHangCauHoi.docx" 
    output_file = "data.js"
    
    convert_docx_to_js(input_file, output_file)