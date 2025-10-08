import docx

def debug_bold_info(docx_path):
    """
    Hàm này chỉ dùng để chẩn đoán. Nó sẽ in ra trạng thái 'bold'
    của từng mẩu văn bản trong các ô đáp án.
    """
    print("--- BẮT ĐẦU PHÂN TÍCH ĐỊNH DẠNG BOLD ---")
    try:
        document = docx.Document(docx_path)
        if not document.tables:
            print("Lỗi: Không tìm thấy bảng nào trong file.")
            return
        
        table = document.tables[0]
        
        # Chỉ kiểm tra 5 câu hỏi đầu tiên cho ngắn gọn
        for i, row in enumerate(table.rows[:5]):
            if len(row.cells) < 3:
                continue
            
            id_text = row.cells[0].text.strip()
            if not id_text.isdigit():
                continue

            print(f"\n[Câu {id_text}]")
            answers_cell = row.cells[2]
            for p in answers_cell.paragraphs:
                if not p.text.strip():
                    continue
                
                print(f"  - Dòng: \"{p.text.strip()}\"")
                
                # Kiểm tra từng 'run' (phần văn bản có cùng định dạng)
                if not p.runs:
                    print("    + Không tìm thấy 'run' nào trong đoạn này.")
                for run in p.runs:
                    # In ra text và trạng thái bold của từng run
                    # run.bold có thể là True (in đậm), False (không in đậm), hoặc None (kế thừa)
                    print(f"    + Run: \"{run.text}\" | run.bold = {run.bold}")

        print("\n--- KẾT THÚC PHÂN TÍCH ---")

    except Exception as e:
        print(f"Lỗi khi đang phân tích file: {e}")

if __name__ == "__main__":
    # Hãy chắc chắn tên file này đúng với file của bạn
    input_file = "NganHangCauHoi.docx"
    
    debug_bold_info(input_file)