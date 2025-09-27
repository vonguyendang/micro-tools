# -*- coding: utf-8 -*-
import os
import sys
import random
import docx
from docx.shared import Pt, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import nsdecls, qn
from docx.oxml import OxmlElement

# --- TÊN CÁC TỆP TIN ĐẦU VÀO ---
QUESTION_BANK_FILE = "NGAN_HANG_DE_GOC.docx"
TEMPLATE_EXAM_FILE = "MAU_DE_THI.docx"
TEMPLATE_ANSWER_FILE = "MAU_DAP_AN.docx"


def parse_question_bank(filename):
    """
    Phân tích ngân hàng câu hỏi, tự động tìm đáp án đúng dựa trên định dạng IN ĐẬM.
    """
    try:
        doc = docx.Document(filename)
        paragraphs = doc.paragraphs
        questions = []
        i = 0
        while i < len(paragraphs):
            text = paragraphs[i].text.strip()
            if text.startswith("Câu "):
                parts = text.split('.', 1)
                q_num_str = parts[0].replace("Câu ", "").strip()
                
                if q_num_str.isdigit() and i + 4 < len(paragraphs):
                    q_num = int(q_num_str)
                    option_paras = paragraphs[i+1 : i+5]
                    correct_option_index = -1

                    for idx, para in enumerate(option_paras):
                        is_correct = any(run.bold for run in para.runs if run.bold)
                        if is_correct:
                            correct_option_index = idx
                            break

                    if correct_option_index != -1:
                        questions.append({
                            "original_q_num": q_num,
                            "question_text": text,
                            "options_text": [p.text.strip() for p in option_paras],
                            "correct_option_index": correct_option_index
                        })
                    else:
                        print(f"⚠️ Cảnh báo: Không tìm thấy đáp án in đậm cho câu {q_num}. Câu này sẽ được bỏ qua.")

                    i += 5
                    continue
            i += 1
        return questions
    except Exception as e:
        print(f"❌ Lỗi khi phân tích ngân hàng đề '{filename}': {e}")
        return []


def generate_shuffled_exam(master_questions, num_questions_per_exam, seed):
    """Trộn thứ tự câu hỏi và thứ tự các lựa chọn đáp án."""
    random.seed(seed)
    shuffled_questions = random.sample(master_questions, num_questions_per_exam)
    
    final_exam_data = []
    new_answer_key = {}

    for new_idx, q_obj in enumerate(shuffled_questions, 1):
        original_correct_idx = q_obj['correct_option_index']
        indexed_options = list(enumerate(q_obj['options_text']))
        random.shuffle(indexed_options)

        new_correct_letter = ''
        for i, (original_idx, text) in enumerate(indexed_options):
            if original_idx == original_correct_idx:
                new_correct_letter = "ABCD"[i]
                break
        
        final_exam_data.append({
            "new_q_num": new_idx,
            "question_text": q_obj['question_text'],
            "shuffled_options": [opt[1] for opt in indexed_options]
        })
        new_answer_key[new_idx] = new_correct_letter
        
    return final_exam_data, new_answer_key


def apply_format_to_paragraph(paragraph, is_bold=False, is_italic=False):
    """Áp dụng font, size, và giãn dòng cho một paragraph."""
    paragraph.paragraph_format.line_spacing = 1.5
    for run in paragraph.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        if is_bold:
            run.bold = True
        if is_italic:
            run.italic = True

def create_exam_file(template_path, output_path, exam_code, exam_data):
    """Tạo file đề thi từ file mẫu, giữ nguyên định dạng."""
    doc = docx.Document(template_path)
    
    for p in doc.paragraphs:
        if "[ ĐỀ GỐC ]" in p.text and "Vĩnh Long" in p.text:
            original_alignment = p.alignment
            p.text = ""
            run_bold = p.add_run(f"[ ĐỀ SỐ {exam_code:02} ]")
            run_bold.bold = True
            p.add_run(" " * 77)
            run_italic = p.add_run("Vĩnh Long, ngày   tháng    năm 2025")
            run_italic.italic = True
            p.alignment = original_alignment
            apply_format_to_paragraph(p)
        elif "[ ĐỀ GỐC ]" in p.text:
            original_alignment = p.alignment
            p.text = ""
            run = p.add_run(f"[ ĐỀ SỐ {exam_code:02} ]")
            run.bold = True
            p.alignment = original_alignment
            apply_format_to_paragraph(p, is_bold=True)
    
    for p in doc.paragraphs:
        if '[[NOI_DUNG_DE]]' in p.text:
            p.text = ""
            for item in exam_data:
                original_text_parts = item['question_text'].split('.', 1)
                question_content = original_text_parts[1].strip() if len(original_text_parts) > 1 else ""
                
                q_para = p.insert_paragraph_before()
                full_question_text = f"Câu {item['new_q_num']}. {question_content}"
                q_para.add_run(full_question_text)
                apply_format_to_paragraph(q_para, is_bold=True)
                
                for i, option_text in enumerate(item['shuffled_options']):
                    opt_letter = f"{chr(65 + i)}."
                    opt_content = option_text.strip()
                    if len(opt_content) > 2 and opt_content[1] == '.' and opt_content[0] in "ABCD":
                        opt_content = opt_content[2:].strip()

                    para_opt = p.insert_paragraph_before()
                    para_opt.add_run(f"    {opt_letter} {opt_content}")
                    apply_format_to_paragraph(para_opt)

                p.insert_paragraph_before()
            break
    doc.save(output_path)


def create_answer_file(template_path, output_path, exam_code, answer_key):
    """Tạo file đáp án từ file mẫu, giữ nguyên định dạng."""
    doc = docx.Document(template_path)
    
    for p in doc.paragraphs:
        if "ĐÁP ÁN ĐỀ THI GỐC" in p.text:
            p.text = ""
            run = p.add_run(f"ĐÁP ÁN ĐỀ THI SỐ {exam_code:02}")
            run.bold = True
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            apply_format_to_paragraph(p)
            break
    
    placeholder = None
    for p in doc.paragraphs:
        if '[[BANG_DAP_AN]]' in p.text:
            placeholder = p
            break
    
    if placeholder:
        num_questions = len(answer_key)
        for i in range(0, num_questions, 20):
            start_q_in_chunk = i + 1
            
            table = doc.add_table(rows=5, cols=21)
            table.style = 'Table Grid'
            
            for row in table.rows: row.height = Cm(0.7)
            table.columns[0].width = Cm(1.2)
            for col_idx in range(1, 21): table.columns[col_idx].width = Cm(0.8)

            hdr_cells = table.rows[0].cells
            
            # ✅ CẬP NHẬT: Thêm chữ CÂU và in đậm
            cell_cau = hdr_cells[0]
            cell_cau.text = ""
            cell_cau.paragraphs[0].add_run("CÂU").bold = True
            
            for j in range(20):
                q_num = start_q_in_chunk + j
                if q_num <= num_questions:
                    cell = hdr_cells[j + 1]
                    cell.text = ""; cell.paragraphs[0].add_run(str(q_num)).bold = True

            for row_idx, letter in enumerate("ABCD", 1):
                row_cells = table.rows[row_idx].cells
                cell = row_cells[0]
                cell.text = ""; cell.paragraphs[0].add_run(letter).bold = True
                
                for j in range(20):
                    q_num = start_q_in_chunk + j
                    if q_num <= num_questions:
                        col_idx = j + 1
                        if answer_key.get(q_num) == letter:
                            shading_elm = OxmlElement('w:shd')
                            shading_elm.set(qn('w:fill'), "D9D9D9")
                            table.cell(row_idx, col_idx)._tc.get_or_add_tcPr().append(shading_elm)
            
            is_last_chunk = (start_q_in_chunk + 19) > num_questions
            if is_last_chunk and (num_questions % 20 != 0):
                merge_start_col_idx = (num_questions % 20) + 1
                for row_idx in range(5):
                    start_cell = table.cell(row_idx, merge_start_col_idx)
                    end_cell = table.cell(row_idx, 20)
                    merged_cell = start_cell.merge(end_cell)
                    set_cell_border(merged_cell, top={'val': 'nil'}, bottom={'val': 'nil'}, start={'val': 'nil'}, end={'val': 'nil'})

            for row in table.rows:
                for cell in row.cells:
                    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                    for paragraph in cell.paragraphs:
                        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        apply_format_to_paragraph(paragraph)
            
            placeholder._p.addprevious(table._tbl)
            spacer = placeholder.insert_paragraph_before("") 
            spacer.paragraph_format.space_before = Pt(6)

        p_element = placeholder._element
        p_element.getparent().remove(p_element)
    
    else:
        print("⚠️ Cảnh báo: Không tìm thấy placeholder [[BANG_DAP_AN]] trong file mẫu đáp án.")

    doc.save(output_path)


def set_cell_border(cell, **kwargs):
    """Hàm tiện ích để ẩn đường viền của một ô."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for border_name, border_attrs in kwargs.items():
        border_el = OxmlElement(f'w:{border_name}')
        for attr_name, attr_val in border_attrs.items():
            border_el.set(qn(f'w:{attr_name}'), str(attr_val))
        tcBorders.append(border_el)
    tcPr.append(tcBorders)

def main():
    """Hàm chính điều khiển luồng của chương trình."""
    print("--- CHƯƠNG TRÌNH TRỘN ĐỀ THI TRẮC NGHIỆM ---")
    
    required_files = [QUESTION_BANK_FILE, TEMPLATE_EXAM_FILE, TEMPLATE_ANSWER_FILE]
    for f in required_files:
        if not os.path.exists(f):
            print(f"❌ Lỗi: Không tìm thấy tệp tin cần thiết '{f}'.")
            print("👉 Vui lòng kiểm tra lại Hướng Dẫn Sử Dụng và chuẩn bị đủ 3 tệp tin.")
            sys.exit(1)
        
    master_questions = parse_question_bank(QUESTION_BANK_FILE)
    if not master_questions:
        print("❌ Phân tích ngân hàng đề thất bại. Vui lòng kiểm tra lại file.")
        sys.exit(1)
    
    total_questions = len(master_questions)
    print(f"✅ Đã phân tích được {total_questions} câu hỏi từ ngân hàng đề.")

    while True:
        try:
            num_exams = int(input("📄 Bạn muốn tạo bao nhiêu mã đề? "))
            if num_exams <= 0: raise ValueError
            break
        except ValueError:
            print("⚠️ Vui lòng nhập một số nguyên dương!")

    while True:
        try:
            prompt = f"✍️  Mỗi đề thi có bao nhiêu câu (tối đa {total_questions})? "
            num_per_exam = int(input(prompt))
            if not (0 < num_per_exam <= total_questions):
                raise ValueError
            break 
        except ValueError:
            print(f"⚠️ Vui lòng nhập một số từ 1 đến {total_questions}!")

    print("\n🚀 Bắt đầu tạo đề...")
    for i in range(1, num_exams + 1):
        exam_data, answer_key = generate_shuffled_exam(master_questions, num_per_exam, seed=i)
        
        exam_filename = f"DE SO {i:02}.docx"
        answer_filename = f"DAP AN DE SO {i:02}.docx"
        
        create_exam_file(TEMPLATE_EXAM_FILE, exam_filename, i, exam_data)
        create_answer_file(TEMPLATE_ANSWER_FILE, answer_filename, i, answer_key)
        
        print(f"✅ Đã tạo thành công: '{exam_filename}' và '{answer_filename}'")
        
    print("\n🎉 HOÀN TẤT! ---")

if __name__ == '__main__':
    main()